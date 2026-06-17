import os
import json
import docx
from pypdf import PdfReader
from groq import Groq
from dotenv import load_dotenv
# Load environment variables
load_dotenv()
def extract_text_from_pdf(file_path):
    """Extract text from a PDF file using pypdf."""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        raise Exception(f"Error reading PDF file: {str(e)}")
    return text
def extract_text_from_docx(file_path):
    """Extract text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {str(e)}")
    return text
def parse_resume_file(uploaded_file):
    """Parse the uploaded file based on its extension."""
    filename = uploaded_file.name
    # Save file temporarily or read bytes directly
    # Since streamlit upload is a BytesIO-like object, we can write a temporary file or pass it to standard tools
    # Let's save it temporarily to read it
    temp_dir = os.path.join(os.path.dirname(__file__), "temp")
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, filename)
    
    try:
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(temp_path)
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(temp_path)
        else:
            # Fallback to plain text
            with open(temp_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    finally:
        # Cleanup
        if os.path.exists(temp_path):
            os.remove(temp_path)
            
    return text
def analyze_resume_ats(resume_text, job_desc_text, api_key=None):
    """
    Sends the resume and job description to Groq Llama 3.3
    and returns a structured JSON payload of the ATS analysis.
    """
    # Use provided key or fall back to environment variable
    key = api_key or os.getenv("GROQ_API_KEY")
    if not key:
        raise ValueError("Groq API Key not found. Please provide it or set GROQ_API_KEY in .env.")
    client = Groq(api_key=key)
    system_prompt = (
        "You are an expert Applicant Tracking System (ATS) algorithm, HR consultant, and technical resume writer. "
        "Your task is to perform a highly thorough, accurate, and professional ATS audit of the user's resume against "
        "the provided job description. You must output the analysis strictly in JSON format. Do not write any conversational "
        "text before or after the JSON payload. Ensure all JSON fields are populated accurately."
    )
    user_prompt = f"""
Analyze the resume and job description below.
--- RESUME ---
{resume_text}
--- JOB DESCRIPTION ---
{job_desc_text}
--- OUTPUT JSON SCHEMA ---
You must return a JSON object with the following exact keys and types:
{{
  "ats_score": <int, 0 to 100 representing general compatibility>,
  "match_percentage": <int, 0 to 100 representing job description keyword/requirement alignment>,
  "formatting_score": <int, 0 to 100 representing layout/font/readability elements>,
  "keyword_score": <int, 0 to 100 representing presence of core resume keywords>,
  "experience_score": <int, 0 to 100 representing relevance of past roles to requirements>,
  "skills_score": <int, 0 to 100 representing technical/soft skills alignment>,
  "analysis_summary": "<string, professional summary of candidate fit, highlighting core alignments and main gaps>",
  "key_strengths": ["<string>", "<string>", ...],
  "critical_issues": ["<string>", "<string>", ...],
  "keyword_analysis": [
    {{
      "keyword": "<string, keyword/skill name>",
      "status": "<string, either 'Present' or 'Missing'>",
      "importance": "<string, either 'High', 'Medium', or 'Low'>"
    }},
    ...
  ],
  "skills_gap": [
    {{
      "skill": "<string, skill/tool name>",
      "category": "<string, e.g. 'Hard Skill', 'Soft Skill', 'Tool/Technology'>",
      "gap_description": "<string, detail of how the requirement is missing or underrepresented in the resume>"
    }},
    ...
  ],
  "ats_compatibility_checks": [
    {{
      "check": "<string, name of standard check, e.g., 'Contact Info Found', 'Standard Headings Used', 'Bullet Points Formatting', 'No Tables/Columns'>",
      "status": "<string, 'Passed', 'Warning', or 'Failed'>",
      "feedback": "<string, constructive feedback on how to fix or why it passed/warned>"
    }},
    ...
  ],
  "improvement_suggestions": ["<string, clear actionable advice>", ...],
  "optimized_resume_snippets": [
    {{
      "original": "<string, sub-optimal bullet point or description in the resume>",
      "optimized": "<string, rewritten action-oriented bullet point that includes keywords and quantifies achievements>",
      "rationale": "<string, why the rewrite is better and what ATS keywords it addresses>"
    }},
    ...
  ],
  "interview_questions": [
    {{
      "question": "<string, tailored behavioral or technical question to ask this candidate during interview based on gaps>",
      "rationale": "<string, why this question is relevant to test the gap>"
    }},
    ...
  ],
  "cover_letter": "<string, a professional, compelling, and fully-formed cover letter matching the job description and leveraging candidate strengths from the resume. Use standard cover letter layout and placeholders for contact fields. Keep formatting clean with newlines.>"
}}
"""
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.2
        )
        
        result_content = response.choices[0].message.content
        return json.loads(result_content)
    except json.JSONDecodeError as je:
        raise Exception(f"Failed to parse AI response as JSON: {str(je)}. Response was: {result_content}")
    except Exception as e:
        raise Exception(f"Error connecting to Groq API: {str(e)}")
